import os
import json
import openai
from pymongo import MongoClient

from docx import Document
import docx2txt as d2t_reader
import fitz as pdf_reader  # PyMuPDF
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import open

from dotenv import dotenv_values


def extract_resume_word(file):
    text = d2t_reader.process(file)
    return text


def extract_resume_pdf(file):
    pdf = pdf_reader.open(file)
    text = "\n".join([page.get_text() for page in pdf])
    return text


def extract_text_from_resume(file):
    if file.endswith(".docx"):
        return extract_resume_word(file)
    elif file.endswith(".pdf"):
        return extract_resume_pdf(file)
    else:
        raise Exception('Invalid resume format. Please upload a valid format')


def add_header_with_image(document, image_path, text):
    header_section = document.sections[0]
    header = header_section.header
    header_text = header.paragraphs[0]
    header_text.text = text
    header_text.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def add_heading(document, text):
    document.add_heading(text, level=1)


def add_paragraph(document, text):
    document.add_paragraph(text)


def add_bullet_list(document, items):
    for item in items:
        document.add_paragraph(item, style='List Bullet')


def add_experience(document, experiences):
    for exp in experiences:
        document.add_paragraph(exp['title'], style='Heading 2')
        # document.add_paragraph(f"{exp['company']}\n{exp['duration']}") puede reemplazarse por duration = datestar y dateend juntos
        document.add_paragraph(f"{exp['company']}\n{exp['date_start']} - {exp['date_end']}")
        document.add_paragraph(exp['description'])


def create_document(data):
    doc = Document()

    add_header_with_image(doc, 'Template/soaint-cv-header.png',
                          f"{data["name"]}\n{data["nationality"]}\n{data["address"]}")

    # Add content
    add_heading(doc, "RESUMEN")
    add_paragraph(doc, data["summary"])

    add_heading(doc, "FORMACIÓN")
    add_bullet_list(doc, data["education"])

    add_heading(doc, "CURSOS Y CERTIFICACIONES")
    add_bullet_list(doc, data["certification"])

    add_heading(doc, "TECNOLOGÍAS")
    add_bullet_list(doc, data["skill"])

    add_heading(doc, "EXPERIENCIA")
    add_experience(doc, data["experience"])

    # Se puede guardar internamente en un repo s3 con el _id para eficiencia de busqueda
    doc.save(f'Template/SOAINT_{data["name"]}.docx')


def convert_resume_info(data):
    config = dotenv_values(".env")
    openai.api_key = config["OPENAI_API_KEY"]

    messages = [
        {
            "role": "system",
            "content": """You will be provided with a resume information of a candidate in large text.
       Your task is to find and extract the required data from resume information and then return a JSON format.
       Here an example for desired JSON Format and explanation of each field:
        {
            "name": <Name and last name of candidate>,
            "nationality": <Nationality or Country of candidate>,
            "address": <Address of candidate>,
            "summary": <Abstract/About me/Profile paragraph of the candidate>,
            "education": <List of degree obtained in schools 
                        (Ex: Bachelor of Science in Computer Science, University of Technology, 2015-2018)>,
            "certification": <List of certificates obtained, courses of the candidate>
            "skill": <List of hard or soft skills/tools of the candidate 
                      (Ex: programming languages, methodologies, tools)>
            "experience": <List of work experience in companies of the candidate. Include actual position>
                        Consider these fields in each element of experience: 
                        (title: <Position in the Job>, company: <Name of the Company>, 
                        date_start: <start date in Month-Year format>, date end: <start date in Month-Year format>,
                        description: <Description of the job/ position, participated or in course projects and achievements>). 
        }
        In case information is not found for any field of the JSON output, set value to "N/A".      
       """,
        },
        {
            "role": "user",
            "content": data,
        }
    ]

    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=messages
    )

    print("*********resultado extraccion**********")
    print(response)

    resume_info = json.loads(response.choices[0].message.content)
    resume_info["full_resume_text"] = data
    print(resume_info)
    return resume_info


def get_summary_ai(data):
    config = dotenv_values(".env")
    openai.api_key = config["OPENAI_API_KEY"]

    messages = [
        {
            "role": "system",
            "content": f"""You will be provided with a resume information of a candidate in large text or json format.
        Your task is to summarize the resume highlighting critical information about specific requirements of positions or projects.
        Your will return the summary in spanish. 
        """,
        },
        {
            "role": "user",
            "content": data,
        }
    ]

    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=messages
    )
    print("*********Resumen**********")
    print(response.choices[0].message.content)
    write_txt_file("output/summary.txt",response.choices[0].message.content)


def get_tech_experience_insight(str_data):
    config = dotenv_values(".env")
    openai.api_key = config["OPENAI_API_KEY"]
    messages = [
        {
            "role": "system",
            "content": """You will be provided with information about candidates resume data in a large text in english or spanish.
            Each candidate resume information is separated by the word ##RESUME##.
            Your task is to analyze the education,skills,experience and certification mentioned in the information of each candidate. Identify the most common programming languages, frameworks, tools, technologies and methodologies. 
            You will highlight first the top skills found among the candidates.
            You will return a structured list with programming languages,frameworks,tools,methodologies and technologies mentioned, 
            along with their frequencies and expertise level for each item.
            For expertise level consider 0-3 where 0 is No Experience and 3 Expert. 
            To determine expertise level take into consideration factors like education,skills, certification and experience in technology 
            Here an example of each element from the list:
            
                    -Tecnología: <Programming Language, framework, tool,methodology or technology> (Example: Python),
                    -Frecuencia: <Frecuency used> (Example: 3),
                    -Experiencia Promedio": <Expertise level of technology> (Example: 3), 
            """,
        },
        {
            "role": "user",
            "content": str_data,
        }
    ]

    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=messages
    )
    print("*********resultado insight 1**********")
    print(response.choices[0].message.content)
    write_txt_file("output/insight_1.txt", response.choices[0].message.content)


def get_featured_clients_projects_insight(str_data):
    config = dotenv_values(".env")
    openai.api_key = config["OPENAI_API_KEY"]
    messages = [
        {
            "role": "system",
            "content": """You will be provided with information about candidates resume data in a large text in english or spanish.
            Each candidate resume information is separated by the word ##RESUME##.
            Your task is to analyze the employment histories of each candidate to identify mentions of work carried out for B2C clients. 
            Extract details such as the type of projects undertaken and the technologies utilized.
            Translate to spanish.
            Here an example for desired output  and explanation of each field:
                        Cliente: (Nombre del cliente Example: Verizon),
                        Proyectos: 
                            - <Project 1 Description>
                            - <Project 2 Description>
                        Tecnologias: <Technology 1, Technology 2, Technology 3>
            
            Important: Don't consider education courses or certifications as a client or project.  
            """,
        },
        {
            "role": "user",
            "content": str_data,
        }
    ]

    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=messages
    )
    print("*********resultado insight 2**********")
    print(response.choices[0].message.content)
    write_txt_file("output/insight_2.txt", response.choices[0].message.content)


def get_competencies_skills_analysis_insight(str_data):
    config = dotenv_values(".env")
    openai.api_key = config["OPENAI_API_KEY"]
    messages = [
        {
            "role": "system",
            "content": """You will be provided with information about candidates resume data in a large text in english or spanish.
            Each candidate resume information is separated by the word ##RESUME##.
            Your task is to analyze the competencies and skills outlined to uncover patterns and emerging trends. 
            Identify areas of strength and potential skill gaps within the technological landscape. 
            This analysis will inform targeted training and development initiatives to enhance organizational capabilities.
            Translate to spanish.
            
            Output: A report in general outlining prevalent competencies, skills and areas for development identified,
            structured to facilitate actionable insights for skill enhancement strategies.
             
            
            """,
        },
        {
            "role": "user",
            "content": str_data,
        }
    ]

    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=messages
    )
    print("*********resultado insight 3**********")
    print(response.choices[0].message.content)
    write_txt_file("output/insight_3.txt", response.choices[0].message.content)

def get_sectorial_experience_insight(str_data):
    config = dotenv_values(".env")
    openai.api_key = config["OPENAI_API_KEY"]
    messages = [
        {
            "role": "system",
            "content": """You will be provided with information about candidates resume data in a large text in english or spanish.
            Each candidate resume information is separated by the word ##RESUME##.
            Your task is to analyze the data employment  and summary histories of candidates.
            Then highlight the experience in different sectors or industries (For example: Banking)
            in order to show the diversity of talent on each sector or industry. 
            Translate to spanish.
            
            The output will be a list containing for each industry or sector:
            industry or sector, number of employers specialized, average time in years of expertise and expertise level
            For expertise level consider 0-3 where 0 is No Experience and 3 Expert. To determine expertise level take into consideration factors like years, different roles, number of employers in sector and achievements.             
            
            Example Output:
            - Sector or Industry: Banking
            - Number of employers: 5
            - Average Experience: 5 years
            - Expertise Level: 2 (Medium)

            """,
        },
        {
            "role": "user",
            "content": str_data,
        }
    ]

    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=messages
    )
    print("*********resultado insight 4**********")
    # print(response)
    print(response.choices[0].message.content)
    write_txt_file("output/insight_4.txt", response.choices[0].message.content)


def save_resume_db(data):
    config = dotenv_values(".env")
    uri = config["MONGO_URI"]
    client = MongoClient(uri)
    db = client[config["MONGO_DATABASE"]]
    collection = db["resumes"]
    collection.insert_one(data)


def get_all_resumes():
    config = dotenv_values(".env")
    uri = config["MONGO_URI"]
    client = MongoClient(uri)
    db = client[config["MONGO_DATABASE"]]
    collection = db["resumes"]
    cursor_result = collection.find()
    list_resumes = list(cursor_result)
    return list_resumes


def build_string_resume_all(list_resumes):
    list_full_text_resume = []
    for item in list_resumes:
        list_full_text_resume.append(item['full_resume_text'])

    all_resumes_str = "##RESUME##".join(str(element) for element in list_full_text_resume)
    return all_resumes_str

def write_txt_file(file,message):
    with open(file, "w+") as archivo:
        archivo.writelines(message)


def main():
    print("***Bienvenidos al programa de conversión de CVs:****")
    while True:
        print("1. Leer CV")
        print("2. Resumen de CV")
        print("3. Conversión de CV")
        print("4. **Analizar: Perfiles de Tecnología y Experiencia**")
        print("5. **Analizar: Clientes y Proyectos Destacados**")
        print("6. **Analizar: Competencias y Habilidades**")
        print("7. **Analizar: Experiencia Sectorial**")
        print("8. Salir")

        choice = input("Por favor, Ingrese una opción (1-8): ")

        if choice == '1':
            print("****LISTADO DE ARCHIVOS EN EL DIRECTORIO DE PROCESAMIENTO***********")
            directory = "CV/"

            for name in os.listdir(directory):
                print(f"Archivo: '{name}'")

            print("!!IMPORTANTE: El archivo debe encontrarse en la ruta de los CVs y ser uno del listado!!!")
            resume_file = input("Por favor, ingrese el nombre del archivo PDF o Word: ")
            is_valid_file = None

            for name in os.listdir(directory):
                if name == resume_file:
                    is_valid_file = True

            if is_valid_file:
                print("***Lectura CV***")
                file_path = os.path.join(directory, resume_file)
                full_resume_text = extract_text_from_resume(file_path)
                json_data = convert_resume_info(full_resume_text)
                print("***Inicio Proceso de guardado de información***")
                save_resume_db(json_data)
                print("!!Enhorabuena!! Información de CV guardada con exito =)")

            else:
                print("Lo sentimos, el archivo no se encuentra en el listado")

        elif choice == '2':
            list_resumes = get_all_resumes()

            if len(list_resumes):
                print("****Listados de CVs Disponibles****")
                for item in list_resumes:
                    print(f"Codigo: {item['_id']} | Talento: {item['name']}")
                resume_code = input("Por favor, Ingrese el código del CV a procesar: ")
                full_resume_text = None
                for item in list_resumes:
                    if str(item['_id']) == str(resume_code):
                        full_resume_text = item['full_resume_text']
                        break

                if full_resume_text:
                    print("***Inicio Generación Resumen***")
                    get_summary_ai(full_resume_text)
                    print("Resumen obtenido con exito")

                else:
                    print("Lo sentimos, no se encontró el código indicado")
            else:
                print("No existen CVs, debe cargar uno primero para generar el resumen")

        elif choice == '3':
            list_resumes = get_all_resumes()

            if len(list_resumes):
                print("****Listados de CVs Disponibles****")
                for item in list_resumes:
                    print(f"Codigo: {item['_id']} | Talento: {item['name']}")
                resume_code = input("Por favor, Ingrese el código del CV a procesar: ")
                resume_data = None
                for item in list_resumes:
                    if str(item['_id']) == str(resume_code):
                        resume_data = item
                        break

                if resume_data:
                    print("*****Creacion de word plantilla******")
                    create_document(resume_data)
                    print("*****Plantilla Creada y guardada con exito*****")

                else:
                    print("Lo sentimos, no se encontró el código indicado")
            else:
                print("No existen CVs, debe cargar uno primero para realizar la conversion")

        elif choice == '4':
            list_resumes = get_all_resumes()
            all_resumes_str = build_string_resume_all(list_resumes)
            print("*****Obteniendo insight: Perfiles de Tecnología y Experiencia*******")
            get_tech_experience_insight(all_resumes_str)
            print("Procesamiento finalizado. Revisar contenido")

        elif choice == '5':
            list_resumes = get_all_resumes()
            all_resumes_str = build_string_resume_all(list_resumes)
            print("*****Obteniendo insight: Clientes y Proyectos Destacados*******")
            get_featured_clients_projects_insight(all_resumes_str)
            print("Procesamiento finalizado. Revisar contenido")

        elif choice == '6':
            list_resumes = get_all_resumes()
            all_resumes_str = build_string_resume_all(list_resumes)
            print("*****Obteniendo insight: Análisis de Competencias y Habilidades*******")
            get_competencies_skills_analysis_insight(all_resumes_str)
            print("Procesamiento finalizado. Revisar contenido")

        elif choice == '7':
            list_resumes = get_all_resumes()
            all_resumes_str = build_string_resume_all(list_resumes)
            print("*****Obteniendo insight: Experiencia Sectorial*******")
            get_sectorial_experience_insight(all_resumes_str)
            print("Procesamiento finalizado. Revisar contenido")

        elif choice == '8':
            print("Gracias por su preferencia. !!Hasta Pronto!!")
            break

        else:
            print("Opción inválida. Ingrese una opción del 1 al 8.")


if __name__ == "__main__":
    main()
