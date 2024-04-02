from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

import uuid


# cv_path = "CV/DOC-20240227-WA0055.docx"
# cv_path = "CV/EDER LARA CV.docx"
# cv_path = "CV/CVHenryValdivia.docx"

# doc = Document(cv_path)

# for paragraph in doc.paragraphs:
# print(paragraph.text)


# import openai
# from dotenv import dotenv_values

# config = dotenv_values(".env")

# color = config["COLOR"]
# openai.api_key = config["OPENAI_API_KEY"]

# print(color)
# print(openai.api_key)

def extract_text_from_file(file):
    # Sample CV text (replace this with your actual CV text)
    cv_text = """
    John Doe
    123 Main St, City, Country
    john.doe@email.com
    Phone: 123-456-7890

    Summary:
    Experienced software engineer with a passion for developing innovative solutions.

    Education:
    Bachelor of Science in Computer Science, University XYZ, Graduated 2020

    Courses and Certifications:
    - Certification in Machine Learning, Coursera
    - Project Management Professional (PMP), PMI

    Skills:
    - Programming: Python, Java, C++
    - Web Development: HTML, CSS, JavaScript
    - Database: SQL, MongoDB

    Experience:
    Software Engineer, ABC Corp, 2020-present
    - Developed and maintained web applications using Django framework.
    - Collaborated with cross-functional teams to deliver projects on time.

    """

    return cv_text


def convert_resume_info(text):
    resume_info = {
        "Nombre y Apellido": "John Doe",
        "Nacionalidad": "American",
        "Ubicación": "New York, USA",
        "Resumen": "Como Ingeniero de Sistemas colegiado desde 2021, con 8 años de experiencia destacada en el sector financiero, he liderado proyectos estratégicos de automatización en áreas críticas como gestión de riesgo operacional, continuidad del negocio y seguridad de la información. Actualmente, me desempeño como líder del equipo de desarrollo de software Normativo y BI en la Caja Municipal de Sullana. Mi trayectoria incluye una participación activa en proyectos de implementación de sistemas web Java, sistemas de escritorio C#, y despliegue de plataformas BI. Me enorgullece haber contribuido al éxito de diversos proyectos gracias a mi sólido conocimiento en lenguajes como Java, JavaScript, Oracle Developer, .NET C#, y frameworks de desarrollo como Spring, Spring MVC y Spring Boot. Además, poseo la certificación internacional de Microsoft AZ-900. Como egresado de la Maestría en Administración de Negocios (MBA), estoy comprometido con el crecimiento personal y profesional, manteniéndome actualizado con las últimas tendencias tecnológicas. Mi enfoque proactivo y dedicación a la mejora continua me preparan para abordar los retos del mañana. Estoy ansioso por contribuir con mi experiencia y habilidades a nuevos proyectos y desafíos",
        "Formación": "- Bachelor's Degree in Computer Science, University XYZ, 2014-2018\n- Master's Degree in Business Administration, University ABC, 2018-2020",
        # "Courses_Certifications": "- Certification in Data Science (Coursera)\n- Project Management Professional (PMP)",
        # "Skills": "- Python\n- Java\n- SQL",
        # "Experience": "- Position: Software Engineer\n  Company: ABC Inc.\n  Date Start: Jan 2018\n  Date End: Present\n\n- Position: Project Manager\n  Company: XYZ Corp.\n  Date Start: Jan 2015\n  Date End: Dec 2017"
    }
    return resume_info


def fill_header_template(doc, info):
    for section in doc.sections:
        header = section.header
        for key, value in info.items():
            for paragraph in header.paragraphs:
                print(f"Key: {key.upper()} , Parragraph: {paragraph.text}")
                if f"{key.upper()}" in paragraph.text.upper():
                    if paragraph.text.isupper():
                        paragraph.text = paragraph.text.replace(f"{key.upper()}", value)
                    else:
                        paragraph.text = paragraph.text.replace(f"{key}", value)


def fill_body_template(doc, info):
    # return "Hola mundo"

    # for paragraph in doc.paragraphs:
    #    print(f"Parragraph: {paragraph.text}")
    #    if paragraph.style.name.startswith('List Paragraph'):
    # Remove the paragraph containing the bullet list
    #        p = paragraph._element
    #        p.getparent().remove(p)

    section_to_replace_next_line = None
    for key, value in info.items():
        for paragraph in doc.paragraphs:

            print(f"Key: {key} , Parragraph: {paragraph.text}")
            print(f"last key value: {section_to_replace_next_line}")

            if section_to_replace_next_line:
                if section_to_replace_next_line == "FORMACIÓN":
                    print("Agregar Lista de formacion")
                    # Insert education items as bullet points
                    education_list = value.split("\n")
                    for item in education_list:
                        if item.strip():
                            doc.add_paragraph(item, style='List Paragraph')
                    section_to_replace_next_line = None
                elif section_to_replace_next_line == "RESUMEN":
                    paragraph.text = paragraph.text.replace(paragraph.text, value)
                    section_to_replace_next_line = None

            if f"{key.upper()}" in paragraph.text:
                section_to_replace_next_line = paragraph.text


def fill_template(template_path, output_path, info):
    doc = Document(template_path)

    fill_header_template(doc, info)
    fill_body_template(doc, info)

    doc.save(output_path)
    print(f"Filled resume saved to: {output_path}")


def main():
    # Input data
    cv_input = "CV/CVHenryValdivia.docx"
    cv_template = "Template/Soaint_Template.docx"

    doc = Document(cv_input)

    for paragraph in doc.paragraphs:
        print(paragraph.text)

    # Extract text from resume
    resume_text = extract_text_from_file(cv_input)

    # Extracted information from the resume into dictinary
    #resume_info = convert_resume_info(resume_text)

    # guid = uuid.uuid1()
    #output_cv = f"Template/soaint-Test.docx"
    #print(output_cv)

    # Create new file
    #fill_template(cv_template, output_cv, resume_info)


if __name__ == "__main__":
    main()
